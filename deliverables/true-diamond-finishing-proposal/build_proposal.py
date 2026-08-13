from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


OUT = "/Users/timkuiper/Documents/Website Autopilots/deliverables/true-diamond-finishing-proposal/True_Diamond_Finish_90_Day_Growth_Pilot.docx"

AP_BG = "F5F5F2"
AP_CARD = "FFFFFF"
AP_SOFT = "F8F8F5"
AP_TEXT = "111111"
AP_MUTED = "656565"
AP_BROWN = "9F3826"
AP_BROWN_DARK = "7D2A1D"
AP_GREEN = "DCE9DD"
AP_LINE = "D9D9D3"
WHITE = "FFFFFF"

BODY_FONT = "Public Sans"
ACCENT_FONT = "Syne"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=AP_LINE, size=8):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_font(run, name=BODY_FONT, size=10.5, color=AP_TEXT, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(p, text, **kwargs):
    r = p.add_run(text)
    set_font(r, **kwargs)
    return r


def add_label(doc, text, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    add_text(p, text.upper(), name=ACCENT_FONT, size=9, color=AP_BROWN, bold=True)
    return p


def add_body(doc, text, after=8, bold_prefix=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        add_text(p, bold_prefix, size=10.5, bold=True)
        add_text(p, text[len(bold_prefix):], size=10.5)
    else:
        add_text(p, text, size=10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_font(r, size=10.25)
    return p


def add_number(doc, title, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_together = True
    add_text(p, title + " — ", size=10.25, bold=True)
    add_text(p, text, size=10.25)
    return p


def add_callout(doc, label, headline, body, fill=AP_SOFT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, AP_LINE, 7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    add_text(p, label.upper(), name=ACCENT_FONT, size=8.5, color=AP_BROWN, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)
    p2.paragraph_format.keep_with_next = True
    add_text(p2, headline, size=13, bold=True)
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.line_spacing = 1.2
    add_text(p3, body, size=10.1, color=AP_MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metric_strip(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [9360 // len(items)] * len(items)
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    set_table_borders(table, AP_LINE, 7)
    for idx, (value, label) in enumerate(items):
        cell = table.cell(0, idx)
        set_cell_shading(cell, AP_CARD if idx != 1 else AP_SOFT)
        set_cell_margins(cell, top=170, start=160, bottom=170, end=160)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        add_text(p, value, name=ACCENT_FONT, size=17, color=AP_BROWN, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        add_text(p2, label, size=8.5, color=AP_MUTED, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def style_table_text(table, header=True):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.12
                for run in p.runs:
                    set_font(
                        run,
                        size=9.25 if r_idx else 9.2,
                        color=WHITE if header and r_idx == 0 else AP_TEXT,
                        bold=header and r_idx == 0,
                    )


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=AP_MUTED)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.82)
sec.right_margin = Inches(0.82)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = BODY_FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(AP_TEXT)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Title", 28, AP_TEXT, 0, 6),
    ("Subtitle", 12.5, AP_MUTED, 0, 8),
    ("Heading 1", 19, AP_TEXT, 16, 8),
    ("Heading 2", 13.5, AP_TEXT, 11, 5),
    ("Heading 3", 11.5, AP_BROWN_DARK, 8, 4),
]:
    st = styles[style_name]
    st.font.name = BODY_FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    st.font.size = Pt(size)
    st.font.bold = style_name != "Subtitle"
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Number"):
    st = styles[style_name]
    st.font.name = BODY_FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    st.font.size = Pt(10.25)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.18

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
add_text(hp, "AUTOPILOTS × SMARTLEADS", name=ACCENT_FONT, size=8.5, color=AP_BROWN, bold=True)
add_text(hp, "  |  TRUE DIAMOND FINISH", size=8.5, color=AP_MUTED, bold=True)

footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_text(fp, "90-DAY GROWTH PILOT   •   ", name=ACCENT_FONT, size=8, color=AP_MUTED, bold=True)
add_page_number(fp)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(36)
add_label(doc, "A practical pilot built around measurable growth", after=14)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
add_text(p, "A local customer engine for", size=28, color=AP_TEXT, bold=True)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(12)
add_text(p2, "True Diamond Finish.", size=28, color=AP_BROWN, bold=True, italic=True)

sub = doc.add_paragraph(style="Subtitle")
sub.paragraph_format.space_after = Pt(22)
add_text(
    sub,
    "A 90-day pilot combining local Meta advertising, a conversion-focused landing page and one clear CRM pipeline.",
    size=12.5,
    color=AP_MUTED,
)

add_metric_strip(doc, [
    ("90 days", "PILOT PERIOD"),
    ("$10/day", "STARTING AD BUDGET"),
    ("$1,000", "COMBINED PILOT FEE"),
])

add_callout(
    doc,
    "The promise",
    "We bring the right local car owners to one simple next step.",
    "They can book directly through Square or leave their details for fast follow-up. Every lead and every closed job is tracked, so the next decision is based on evidence—not promises.",
    AP_GREEN,
)

doc.add_paragraph().paragraph_format.space_after = Pt(16)
meta = doc.add_table(rows=3, cols=2)
set_table_geometry(meta, [1900, 7460])
set_table_borders(meta, AP_LINE, 6)
for i, (label, value) in enumerate([
    ("Prepared for", "True Diamond Finish (CelMaura Inc.)"),
    ("Prepared by", "Autopilots & Smartleads"),
    ("Date", "July 30, 2026"),
]):
    set_cell_shading(meta.cell(i, 0), AP_SOFT)
    set_cell_shading(meta.cell(i, 1), AP_CARD)
    meta.cell(i, 0).text = label
    meta.cell(i, 1).text = value
style_table_text(meta, header=False)
for row in meta.rows:
    for run in row.cells[0].paragraphs[0].runs:
        set_font(run, size=9, color=AP_BROWN_DARK, bold=True)

doc.add_page_break()

# 1. Why this pilot
doc.add_heading("1. Why this pilot", level=1)
add_body(
    doc,
    "True Diamond Finish already has the hardest ingredients: a strong mobile service, a large library of before-and-after results, real customer proof and a team that knows how to close. The missing piece is a predictable route from local attention to a booked appointment.",
)
add_body(
    doc,
    "Today, too much growth still depends on door-to-door work, manual follow-up and a broad website experience. The current website shows the full service range and already supports Square booking, but paid traffic needs a faster and more focused message: we come to you, we clean your car at your location, and booking takes only a few steps.",
)
add_callout(
    doc,
    "Pilot focus",
    "One local offer. One focused page. One measurable pipeline.",
    "The pilot is intentionally small. It is designed to learn which audience, creative and offer produce qualified appointments within a practical service radius.",
)

doc.add_heading("The customer journey", level=2)
add_number(doc, "See the ad", "A local car owner sees a clear before-and-after ad within the selected service area.")
add_number(doc, "Choose a next step", "The landing page offers two routes: book directly in Square or leave contact details.")
add_number(doc, "Get followed up", "New enquiries enter the CRM immediately and the True Diamond Finish team receives a clear action.")
add_number(doc, "Track the outcome", "Each lead is moved through the pipeline until booked, won or lost—with the reason recorded.")

doc.add_heading("What success looks like", level=2)
add_bullet(doc, "A consistent source of local enquiries that does not depend on door-to-door availability.")
add_bullet(doc, "A landing page that leads with mobile convenience and local relevance.")
add_bullet(doc, "A visible pipeline from new lead to booked appointment and closed job.")
add_bullet(doc, "Enough clean data to decide whether to scale, adjust the offer or stop.")

# 2. Scope
doc.add_page_break()
doc.add_heading("2. What we will build and run", level=1)

scope = doc.add_table(rows=1, cols=3)
scope.rows[0].cells[0].text = "Workstream"
scope.rows[0].cells[1].text = "Owner"
scope.rows[0].cells[2].text = "Included in the pilot"
rows = [
    ("Landing page", "Autopilots", "Mobile-first page focused on “we come to you,” with local copy, proof, before/after content, direct Square booking and a lead form."),
    ("CRM & pipeline", "Autopilots", "Lead capture, source tracking, pipeline stages, owner assignment, status updates and a simple pilot dashboard."),
    ("Meta campaign", "Smartleads", "Local campaign setup, max. ±20 km radius, audience setup, copy and creative variations, testing and optimization."),
    ("Creative input", "True Diamond Finish", "Best before/after photos, available videos, testimonials, service details, pricing and brand assets."),
    ("Lead handling", "True Diamond Finish", "Fast contact, qualification, appointment confirmation, pipeline updates and accurate closed-job values."),
]
for workstream, owner, included in rows:
    cells = scope.add_row().cells
    cells[0].text = workstream
    cells[1].text = owner
    cells[2].text = included
set_table_geometry(scope, [1700, 1700, 5960])
set_table_borders(scope, AP_LINE, 7)
for c in scope.rows[0].cells:
    set_cell_shading(c, AP_TEXT)
for i, row in enumerate(scope.rows[1:], start=1):
    fill = AP_CARD if i % 2 else AP_SOFT
    for c in row.cells:
        set_cell_shading(c, fill)
style_table_text(scope)

doc.add_heading("Landing-page message", level=2)
add_callout(
    doc,
    "Recommended hero",
    "We come to you. Your car leaves looking like new.",
    "Mobile detailing at your home or workplace in the selected local service area. Choose your service and book directly—or leave your details and we will help you find the right option.",
)
add_bullet(doc, "Primary action: Book your mobile detail.")
add_bullet(doc, "Secondary action: Get help choosing a package.")
add_bullet(doc, "Proof: real transformations, customer testimonials and the strongest three packages.")
add_bullet(doc, "Trust: clear service radius, response expectation, phone number and Square booking.")

doc.add_heading("CRM pipeline", level=2)
add_body(doc, "New lead → Contact attempted → Qualified → Appointment booked → Service completed → Won / Lost")
add_body(
    doc,
    "Each record should include campaign, ad, service interest, booking status, estimated value, actual sale value and lost reason. This creates the evidence needed for the post-pilot pricing model.",
)

# 3. Commercials
doc.add_page_break()
doc.add_heading("3. Pilot investment and commercial terms", level=1)

fees = doc.add_table(rows=1, cols=4)
for idx, text in enumerate(("Item", "Amount", "Timing", "Paid to")):
    fees.rows[0].cells[idx].text = text
fee_rows = [
    ("Pilot fee — first instalment", "$500", "At approval", "Autopilots & Smartleads"),
    ("Pilot fee — final instalment", "$500", "At the end of month 3", "Autopilots & Smartleads"),
    ("Meta media budget", "Approx. $10/day", "Charged directly by Meta", "Meta"),
    ("Estimated 90-day media spend", "Approx. $900", "Based on 90 days", "Meta"),
]
for row_data in fee_rows:
    cells = fees.add_row().cells
    for idx, text in enumerate(row_data):
        cells[idx].text = text
set_table_geometry(fees, [3300, 1600, 2560, 1900])
set_table_borders(fees, AP_LINE, 7)
for c in fees.rows[0].cells:
    set_cell_shading(c, AP_TEXT)
for i, row in enumerate(fees.rows[1:], start=1):
    for c in row.cells:
        set_cell_shading(c, AP_CARD if i % 2 else AP_SOFT)
style_table_text(fees)

add_body(
    doc,
    "The media budget is separate from the pilot fee and remains under True Diamond Finish’s control. Any third-party usage costs outside the items above—such as SMS or additional software licenses—must be approved before activation.",
    after=10,
)

doc.add_heading("The honest break-even view", level=2)
add_metric_strip(doc, [
    ("$150", "AVERAGE REVENUE / JOB"),
    ("$55", "CONTRIBUTION / JOB"),
    ("20%", "ASSUMED CLOSE RATE"),
])
add_body(
    doc,
    "Using the figures discussed, 23 additional customers create about $1,265 in contribution. That is a useful commercial milestone and roughly covers the $1,000 pilot fee plus about $265 in media.",
)
add_body(
    doc,
    "Full 90-day payback is a higher bar. A $1,000 pilot fee plus approximately $900 in media equals about $1,900. At $55 contribution per customer, full break-even requires approximately 35 incremental customers.",
)

math = doc.add_table(rows=1, cols=4)
for idx, text in enumerate(("Scenario", "Customers", "Revenue", "Contribution")):
    math.rows[0].cells[idx].text = text
for row_data in [
    ("Commercial milestone", "23", "$3,450", "$1,265"),
    ("Full pilot break-even", "35", "$5,250", "$1,925"),
    ("At $10 CPL + 20% close", "18", "$2,700", "$990"),
]:
    cells = math.add_row().cells
    for idx, text in enumerate(row_data):
        cells[idx].text = text
set_table_geometry(math, [3400, 1700, 2100, 2160])
set_table_borders(math, AP_LINE, 7)
for c in math.rows[0].cells:
    set_cell_shading(c, AP_TEXT)
for i, row in enumerate(math.rows[1:], start=1):
    for c in row.cells:
        set_cell_shading(c, AP_GREEN if i == 2 else (AP_CARD if i % 2 else AP_SOFT))
style_table_text(math)

add_callout(
    doc,
    "Important",
    "A $10 CPL is a guardrail—not a profitability guarantee.",
    "At $900 media spend, a $10 CPL produces about 90 leads. At a 20% close rate, that is 18 customers. To reach 35 customers with the same budget, the pilot would need roughly a $5.15 CPL, a close rate near 39%, a higher average margin, repeat purchases, upsells—or a combination of these.",
    AP_GREEN,
)

# 4. measurement and ladder
doc.add_page_break()
doc.add_heading("4. Measurement and the win-win growth ladder", level=1)
add_body(
    doc,
    "The first 90 days remain simple: the fixed $1,000 pilot fee covers the build, campaign management and learning period. There is no additional success fee during the pilot. After 90 days, both sides review the same CRM data and choose the next model.",
)

doc.add_heading("Pilot scorecard", level=2)
score = doc.add_table(rows=1, cols=4)
for idx, text in enumerate(("Metric", "Target / guardrail", "Why it matters", "Source")):
    score.rows[0].cells[idx].text = text
score_rows = [
    ("Cost per lead", "≤ $10 target", "Controls acquisition efficiency", "Meta + CRM"),
    ("Lead response time", "≤ 10 minutes in business hours", "Improves the chance of booking", "CRM"),
    ("Lead-to-customer rate", "20% baseline", "Tests sales follow-up quality", "CRM"),
    ("Booked appointment rate", "Measured weekly", "Separates lead quality from closing", "Square + CRM"),
    ("Contribution per job", "$55 baseline", "Connects marketing to real economics", "TDF"),
    ("Attributable customers", "23 milestone / 35 break-even", "Shows commercial and full-payback progress", "CRM"),
]
for row_data in score_rows:
    cells = score.add_row().cells
    for idx, text in enumerate(row_data):
        cells[idx].text = text
set_table_geometry(score, [2100, 2100, 3560, 1600])
set_table_borders(score, AP_LINE, 7)
for c in score.rows[0].cells:
    set_cell_shading(c, AP_TEXT)
for i, row in enumerate(score.rows[1:], start=1):
    for c in row.cells:
        set_cell_shading(c, AP_CARD if i % 2 else AP_SOFT)
style_table_text(score)

doc.add_heading("Recommended post-pilot ladder", level=2)
add_body(
    doc,
    "The next agreement should contain a modest base fee for the real monthly work, plus a success fee only after True Diamond Finish has passed the agreed contribution hurdle. That prevents the partners from being rewarded while the client is still underwater.",
)

ladder = doc.add_table(rows=1, cols=4)
for idx, text in enumerate(("90-day attributable result", "Success fee", "Meaning", "Next move")):
    ladder.rows[0].cells[idx].text = text
ladder_rows = [
    ("0–22 completed jobs", "$0 per job", "Pilot has not reached the commercial milestone", "Fix offer, creative, page or follow-up before scaling"),
    ("23–34 completed jobs", "$0 per job", "Strong signal, but full pilot cost is not yet recovered", "Maintain budget and improve conversion"),
    ("35–45 completed jobs", "$10 per job above 35", "Full pilot economics have crossed break-even", "Scale carefully"),
    ("46–60 completed jobs", "$15 per job above 45", "Profitable, repeatable acquisition is emerging", "Increase media in controlled steps"),
    ("61+ completed jobs", "$20 per job above 60", "High-volume performance with room to share upside", "Agree capacity and expand geography/offers"),
]
for row_data in ladder_rows:
    cells = ladder.add_row().cells
    for idx, text in enumerate(row_data):
        cells[idx].text = text
set_table_geometry(ladder, [2500, 1900, 2900, 2060])
set_table_borders(ladder, AP_LINE, 7)
for c in ladder.rows[0].cells:
    set_cell_shading(c, AP_TEXT)
for i, row in enumerate(ladder.rows[1:], start=1):
    for c in row.cells:
        set_cell_shading(c, AP_GREEN if i >= 3 else (AP_CARD if i % 2 else AP_SOFT))
style_table_text(ladder)

add_body(
    doc,
    "Final post-pilot base fee, attribution window and tier thresholds will be confirmed after the first 90 days. A suggested attribution rule is: a completed and paid job is attributable when the original lead entered through the pilot campaign or landing page within the prior 30 days.",
    after=6,
)
add_body(
    doc,
    "Why this is win-win: True Diamond Finish keeps the first part of the upside to recover the pilot investment. Autopilots and Smartleads participate only when additional completed jobs create measurable value.",
)

# 5. timeline, responsibilities
doc.add_heading("5. 90-day rollout", level=1)
timeline = doc.add_table(rows=1, cols=3)
for idx, text in enumerate(("Period", "What happens", "Decision / output")):
    timeline.rows[0].cells[idx].text = text
timeline_rows = [
    ("Week 1", "Offer workshop, service radius, access, assets, tracking plan and CRM stages.", "Pilot brief approved"),
    ("Week 2", "Landing page build, Square route, lead form, CRM pipeline and campaign preparation.", "Ready for review"),
    ("Week 3", "Quality check, mobile review, tracking test and campaign launch.", "Pilot goes live"),
    ("Weeks 4–6", "Test before/after angles, copy and audiences. Review lead quality and response speed.", "First optimization cycle"),
    ("Weeks 7–10", "Keep winners, replace weak creatives and improve booking friction.", "Second optimization cycle"),
    ("Weeks 11–13", "Consolidate results, verify closed revenue and calculate unit economics.", "90-day review and next-step proposal"),
]
for row_data in timeline_rows:
    cells = timeline.add_row().cells
    for idx, text in enumerate(row_data):
        cells[idx].text = text
set_table_geometry(timeline, [1600, 5200, 2560])
set_table_borders(timeline, AP_LINE, 7)
for c in timeline.rows[0].cells:
    set_cell_shading(c, AP_TEXT)
for i, row in enumerate(timeline.rows[1:], start=1):
    for c in row.cells:
        set_cell_shading(c, AP_CARD if i % 2 else AP_SOFT)
style_table_text(timeline)

doc.add_page_break()
doc.add_heading("What we need from True Diamond Finish", level=2)
add_bullet(doc, "One decision-maker for the pilot and one person responsible for daily lead follow-up.")
add_bullet(doc, "Access needed for Meta, Square booking, domain/landing-page setup and CRM users.")
add_bullet(doc, "A curated asset folder: strongest 15–25 before/after sets, 3–5 videos and 3–5 testimonials.")
add_bullet(doc, "Confirmed packages, starting prices, service radius, working hours and appointment capacity.")
add_bullet(doc, "Pipeline updates within 48 hours and actual job value after completion.")
add_bullet(doc, "A practical response commitment during business hours; slow follow-up will be visible in the reporting.")

doc.add_heading("Boundaries", level=2)
add_bullet(doc, "No guaranteed number of leads, bookings or sales. Paid media performance depends on market response, creative, seasonality, capacity and follow-up.")
add_bullet(doc, "The initial campaign is B2C mobile detailing within one agreed local service area. B2B dealership outreach and AI voice follow-up are separate future workstreams.")
add_bullet(doc, "Major website rebuilds, new content production, advanced AI voice agents and outbound sales services are not included in this pilot.")
add_bullet(doc, "True Diamond Finish owns its customer data, ad account and payment relationship with Meta.")

# 6 acceptance
doc.add_page_break()
doc.add_heading("6. Decision", level=1)
add_callout(
    doc,
    "Recommended next step",
    "Approve the 90-day pilot and start with one tightly defined local offer.",
    "The first goal is not to scale immediately. It is to prove a repeatable route from a local ad to a completed, profitable mobile-detailing job.",
    AP_GREEN,
)

doc.add_heading("Pilot approval", level=2)
approval = doc.add_table(rows=6, cols=2)
approval_data = [
    ("Pilot", "90-day local growth pilot"),
    ("Pilot fee", "$1,000 total — $500 at approval and $500 at the end of month 3"),
    ("Media", "Approx. $10/day, paid directly to Meta"),
    ("Target launch", "Within approximately 15 business days after access and assets are received"),
    ("True Diamond Finish", "Name / signature / date: __________________________________________"),
    ("Autopilots & Smartleads", "Name / signature / date: __________________________________________"),
]
for i, (label, value) in enumerate(approval_data):
    approval.cell(i, 0).text = label
    approval.cell(i, 1).text = value
    set_cell_shading(approval.cell(i, 0), AP_SOFT)
    set_cell_shading(approval.cell(i, 1), AP_CARD)
set_table_geometry(approval, [2200, 7160])
set_table_borders(approval, AP_LINE, 7)
style_table_text(approval, header=False)
for row in approval.rows:
    for run in row.cells[0].paragraphs[0].runs:
        set_font(run, size=9.25, color=AP_BROWN_DARK, bold=True)

doc.add_paragraph().paragraph_format.space_after = Pt(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
add_text(p, "Build the system. Measure the truth. Scale what works.", size=13, bold=True)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_text(p2, "Autopilots × Smartleads", name=ACCENT_FONT, size=9, color=AP_BROWN, bold=True)

doc.core_properties.title = "True Diamond Finish — 90-Day Growth Pilot"
doc.core_properties.subject = "Local Meta advertising, landing page and CRM pilot"
doc.core_properties.author = "Autopilots & Smartleads"
doc.core_properties.keywords = "True Diamond Finish, Autopilots, Smartleads, Meta Ads, CRM, landing page"

doc.save(OUT)
print(OUT)
